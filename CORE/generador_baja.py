from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document


# ============================================================
# API PRINCIPAL
# ============================================================

def generar_solicitud_baja(
    datos: dict[str, Any],
    ruta_docx: Path,
    ruta_odt: Path | None = None,
) -> None:
    """
    Genera la solicitud de baja usando una plantilla DOCX.

    La plantilla debe estar en:
    templates/solicitud_baja_plantilla.docx

    ruta_odt queda como argumento opcional para mantener compatibilidad,
    pero en esta versión no se genera ODT.
    """

    ruta_docx = Path(ruta_docx)
    ruta_docx.parent.mkdir(parents=True, exist_ok=True)

    ruta_plantilla = _obtener_ruta_plantilla()

    if not ruta_plantilla.exists():
        raise FileNotFoundError(
            f"No se encontró la plantilla de baja en: {ruta_plantilla}"
        )

    documento = Document(ruta_plantilla)

    reemplazos = _crear_reemplazos(datos)

    _reemplazar_marcadores_documento(documento, reemplazos)

    documento.save(ruta_docx)


# ============================================================
# RUTAS
# ============================================================

def _obtener_ruta_plantilla() -> Path:
    base_dir = Path(__file__).resolve().parent.parent
    return base_dir / "templates" / "solicitud_baja_plantilla.docx"


# ============================================================
# REEMPLAZOS
# ============================================================

def _texto(valor: Any, fallback: str = "") -> str:
    if valor is None:
        return fallback

    texto = str(valor).strip()
    return texto if texto else fallback


def _tachar_texto(texto: str) -> str:
    """
    Tacha texto usando el carácter Unicode de tachado combinado.
    Evita tener que manipular formato interno de Word.
    """
    return "".join(caracter + "\u0336" for caracter in texto)


def _crear_opcion_tachada(tipo_beneficiario: str) -> str:
    tipo = _texto(tipo_beneficiario)

    if tipo == "Destinatario":
        return f"Destinatario / {_tachar_texto('Tutor')}"

    if tipo == "Tutor":
        return f"{_tachar_texto('Destinatario')} / Tutor"

    return "Destinatario / Tutor"


def _crear_opcion_responsable_informado(valor: str) -> str:
    informado = _texto(valor)

    if informado in {"Sí", "Si"}:
        return "Sí"

    if informado == "No":
        return "No"

    return informado


def _crear_reemplazos(datos: dict[str, Any]) -> dict[str, str]:
    motivo = _texto(datos.get("bajaMotivo"))

    motivos = [
        "Encontrarse privado de la libertad",
        "Fallecimiento",
        "Haber cumplimentado con todos los acuerdos para su egreso",
        "Liquidaciones consecutivas impagas",
        "Mudanza a otro Municipio",
        "Negativa a cumplir con su Acuerdo de Compromiso",
        "Negativa del Joven a participar",
        "Negativa o dificultades al momento de la socialización",
        "Pase de Destinatario a Tutor",
        "Pase de Tutor a Destinatario",
        "Trabajo Formal",
        "Otros motivos",
    ]

    reemplazos = {
        "{{NOMBRE_SEDE}}": _crear_nombre_sede(datos),
        "{{TIPO_MODALIDAD}}": _texto(datos.get("modalidad")),
        "{{NOMBRE_PROFESIONAL}}": _texto(datos.get("profesionalesIntervinientesTexto")),
        "{{NOMBRE_COORDINADORA}}": _texto(datos.get("coordinadora")),
        "{{FECHA_ACTUAL}}": _texto(datos.get("fechaBaja")),

        "{{OPCION_TACHADA}}": _crear_opcion_tachada(
            _texto(datos.get("bajaTipoBeneficiario"))
        ),

        "{{NOMBRE_COMPLETO_BENEFICIARIO}}": _texto(
            datos.get("bajaNombreCompleto")
        ),
        "{{NRO_DNI}}": _crear_tipo_y_dni(datos),
        "{{FECHA_INGRESO}}": _texto(datos.get("bajaFechaIngreso")),

        "{{NOMBRE_COMPLETO_RESPONSABLE_ADULTO}}": _texto(
            datos.get("bajaResponsableNombreCompleto")
        ),
        "{{RELACION}}": _texto(datos.get("bajaResponsableRelacion")),
        "{{OPCION}}": _crear_opcion_responsable_informado(
            _texto(datos.get("bajaResponsableInformado"))
        ),
    }

    for indice, motivo_actual in enumerate(motivos, start=1):
        reemplazos[f"{{{{marca_{indice}}}}}"] = "X" if motivo == motivo_actual else ""

    return reemplazos


def _crear_nombre_sede(datos: dict[str, Any]) -> str:
    centro = _texto(datos.get("centroEnvion"))
    sede = _texto(datos.get("sede"))

    if centro and sede:
        return f"{centro}: {sede}"

    if sede:
        return sede

    return centro


def _crear_tipo_y_dni(datos: dict[str, Any]) -> str:
    tipo_dni = _texto(datos.get("bajaTipoDni"), "DNI")
    dni = _texto(datos.get("bajaDni"))

    if tipo_dni and dni:
        return f"{tipo_dni} {dni}"

    return dni or tipo_dni


# ============================================================
# REEMPLAZO EN DOCUMENTO
# ============================================================

def _reemplazar_marcadores_documento(
    documento: Document,
    reemplazos: dict[str, str],
) -> None:
    """
    Reemplaza marcadores en:
    - cuerpo principal;
    - tablas del cuerpo;
    - encabezados;
    - pies de página.
    """

    # Cuerpo principal
    _reemplazar_marcadores_en_contenedor(documento, reemplazos)

    # Encabezados y pies de página
    for section in documento.sections:
        _reemplazar_marcadores_en_contenedor(section.header, reemplazos)
        _reemplazar_marcadores_en_contenedor(section.footer, reemplazos)


def _reemplazar_marcadores_en_contenedor(
    contenedor,
    reemplazos: dict[str, str],
) -> None:
    for parrafo in contenedor.paragraphs:
        _reemplazar_marcadores_en_parrafo(parrafo, reemplazos)

    for tabla in contenedor.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    _reemplazar_marcadores_en_parrafo(parrafo, reemplazos)


def _reemplazar_marcadores_en_parrafo(
    parrafo,
    reemplazos: dict[str, str],
) -> None:
    """
    Reemplazo simple conservando el formato del primer run del párrafo.

    Para esta plantilla alcanza porque los marcadores están escritos completos.
    Si Google Docs parte internamente algún marcador en varios runs,
    este método reconstruye el texto completo del párrafo.
    """

    texto_original = "".join(run.text for run in parrafo.runs)

    if not texto_original:
        return

    texto_nuevo = texto_original

    for marcador, valor in reemplazos.items():
        texto_nuevo = texto_nuevo.replace(marcador, valor)

    if texto_nuevo == texto_original:
        return

    # Conserva el formato base del primer run y limpia el resto.
    if parrafo.runs:
        parrafo.runs[0].text = texto_nuevo

        for run in parrafo.runs[1:]:
            run.text = ""



